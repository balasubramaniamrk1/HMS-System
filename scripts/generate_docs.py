import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('HealthPlus Management System (HMS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Product Documentation & System Architecture')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n' * 5)
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run('Version: 1.0\n').bold = True
    details.add_run('Date: December 24, 2025\n')
    details.add_run('Prepared for: HealthPlus Hospital')

    doc.add_page_break()

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The HealthPlus Management System (HMS) is a comprehensive web-based platform designed to streamline "
        "hospital operations. It integrates clinical management, inventory control, staff administration, "
        "and public-facing services into a unified ecosystem."
    )

    # --- User Roles ---
    doc.add_heading('2. User Roles & Access Control', level=1)
    doc.add_paragraph(
        "The system employs Role-Based Access Control (RBAC) to ensure security and data privacy. "
        "The following roles are defined:"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Access Scope'
    hdr_cells[2].text = 'Key Responsibilities'
    
    roles = [
        ('Admin', 'Full System', 'User management, configuration, oversight.'),
        ('Doctor', 'Clinical Console', 'Consultations, Prescriptions, Viewing Reports.'),
        ('Inventory Manager', 'Inventory Module', 'Asset tracking, Procurement, AMC Management.'),
        ('Staff', 'Front Desk', 'Appointments, Patient Registration, Attendance.'),
        ('Public User', 'Website', 'Booking appointments, Viewing Information.')
    ]

    for role, scope, resp in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = scope
        row_cells[2].text = resp

    # --- Functional Modules ---
    doc.add_heading('3. Functional Modules', level=1)

    # Inventory
    doc.add_heading('3.1 Inventory Management', level=2)
    doc.add_paragraph("A robust module for tracking high-value assets and consumables.")
    inv_list = doc.add_paragraph(style='List Bullet')
    inv_list.add_run("Equipment Tracking: ").bold = True
    inv_list.add_run("Manage serial numbers, warranty, and status.")
    
    amc_list = doc.add_paragraph(style='List Bullet')
    amc_list.add_run("AMC Management: ").bold = True
    amc_list.add_run("Track maintenance contracts, receive expiry alerts.")
    
    qr_list = doc.add_paragraph(style='List Bullet')
    qr_list.add_run("QR Code Integration: ").bold = True
    qr_list.add_run("Auto-generated QR codes for quick asset identification.")

    # Clinical
    doc.add_heading('3.2 Clinical Management', level=2)
    doc.add_paragraph("Tools for doctors to manage patient care efficiently.")
    
    doc.add_paragraph("Doctor Console features:", style='List Bullet')
    doc.add_paragraph("Digital Appointment Schedule", style='List Bullet 2')
    doc.add_paragraph("Prescription Designer (Auto-generates PDF)", style='List Bullet 2')
    doc.add_paragraph("Patient History View", style='List Bullet 2')

    # --- Database Schema & Mermaid ERD ---
    doc.add_heading('4. Database Schema & ER Diagram', level=1)
    doc.add_paragraph("Below is the detailed schema for all system modules.")

    # Schema Definition (Hardcoded for documentation accuracy)
    modules = {
        "Inventory Management": [
            ("Vendor", [
                ("name", "Char"), ("rating", "Int"), ("phone", "Char"), ("email", "Email")
            ]),
            ("EquipmentCategory", [
                ("name", "Char"), ("description", "Text")
            ]),
            ("Equipment", [
                ("name", "Char"), ("serial_number", "Char (Unique)"), ("status", "Choice"),
                ("vendor", "FK -> Vendor"), ("category", "FK -> EquipmentCategory"),
                ("purchase_date", "Date"), ("warranty_expiry", "Date"),
                ("cost", "Decimal"), ("useful_life", "Int"), ("qr_code", "Image")
            ]),
            ("MaintenanceContract", [
                ("equipment", "FK -> Equipment"), ("vendor", "FK -> Vendor"),
                ("contract_start", "Date"), ("contract_end", "Date"),
                ("cost", "Decimal"), ("is_active", "Calculated")
            ]),
            ("Consumable", [
                ("name", "Char"), ("quantity_in_stock", "Int"),
                ("minimum_stock_level", "Int"), ("vendor", "FK -> Vendor")
            ]),
            ("MaintenanceLog", [
                ("equipment", "FK -> Equipment"), ("date", "Date"),
                ("performed_by", "Char"), ("description", "Text"), ("cost", "Decimal")
            ])
        ],
        "Clinical Management": [
            ("Doctor", [
                ("user", "OneToOne -> User"), ("name", "Char"),
                ("department", "FK -> Department"), ("specialization", "Char")
            ]),
            ("Department", [
                ("name", "Char"), ("slug", "Slug")
            ]),
            ("AppointmentRequest", [
                ("name", "Char"), ("phone", "Char"), ("doctor", "FK -> Doctor"),
                ("preferred_date", "Date"), ("status", "Choice (pending, confirmed...)")
            ]),
            ("Consultation", [
                ("appointment", "OneToOne -> AppointmentRequest"),
                ("diagnosis", "Text"), ("doctor_notes", "Text"),
                ("prescriptions", "Text (Legacy)")
            ]),
            ("PrescriptionItem", [
                ("consultation", "FK -> Consultation"),
                ("medicine_name", "Char"), ("dosage", "Char"), ("duration", "Char")
            ])
        ],
        "Staff Management": [
            ("StaffProfile", [
                ("user", "OneToOne -> User"), ("employee_id", "Char"),
                ("designation", "Char"), ("joining_date", "Date")
            ]),
            ("Attendance", [
                ("staff", "FK -> StaffProfile"), ("date", "Date"),
                ("check_in", "Time"), ("check_out", "Time"), ("status", "Choice")
            ])
        ],
        "Reputation Management": [
            ("Platform", [("name", "Char (Google, Facebook...)")]),
            ("Review", [
                ("platform", "FK -> Platform"), ("author_name", "Char"),
                ("rating", "Int"), ("content", "Text")
            ])
        ]
    }

    # 1. Generate Tables in Doc
    for module_name, tables in modules.items():
        doc.add_heading(f'Module: {module_name}', level=2)
        
        for table_name, fields in tables:
            doc.add_heading(f'Table: {table_name}', level=3)
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Light Shading'
            t.rows[0].cells[0].text = 'Field / Relationship'
            t.rows[0].cells[1].text = 'Type / Details'
            
            for field, type_desc in fields:
                row = t.add_row().cells
                row[0].text = field
                row[1].text = type_desc
            
            doc.add_paragraph() # Spacer

    # 2. Generate Mermaid ER Diagram
    doc.add_heading('5. Deployment & ER Diagram', level=1)
    doc.add_paragraph("The following is the Mermaid.js script to generate the Entity Relationship Diagram (ERD).")
    
    mermaid_code = ["erDiagram"]
    
    # Process relationships for Mermaid
    # Simple logic to find FKs and simple fields
    for module_name, tables in modules.items():
        for table_name, fields in tables:
            mermaid_code.append(f"    {table_name} {{")
            for field, type_desc in fields:
                # Add relationship lines
                if "FK ->" in type_desc:
                    target = type_desc.split("->")[1].strip()
                    mermaid_code.insert(1, f"    {table_name} }}|..|| {target} : links_to")
                    clean_type = "FK"
                elif "OneToOne ->" in type_desc:
                    target = type_desc.split("->")[1].strip()
                    mermaid_code.insert(1, f"    {table_name} ||--|| {target} : matches")
                    clean_type = "OneToOne"
                else:
                    clean_type = type_desc.split(' ')[0]
                
                mermaid_code.append(f"        {clean_type} {field}")
            mermaid_code.append("    }")

    final_mermaid = "\n".join(mermaid_code)
    
    # Add to Docx
    p = doc.add_paragraph()
    runner = p.add_run(final_mermaid)
    runner.font.name = 'Courier New'
    runner.font.size = Pt(9)
    
    # Save Mermaid file separately
    with open('hms_er_diagram.mmd', 'w') as f:
        f.write(final_mermaid)

    doc.save('HMS_Product_Documentation.docx')
    print("Document generated: HMS_Product_Documentation.docx")
    print("Mermaid Diagram generated: hms_er_diagram.mmd")

if __name__ == "__main__":
    create_document()
